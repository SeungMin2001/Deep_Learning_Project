from dotenv import load_dotenv
import os
import re
from langchain_community.chat_models import ChatOpenAI
from langchain.prompts import ChatPromptTemplate
from tabulate import tabulate  # pip install tabulate
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

class Step5:
    def last(self,x):
        import os  # 함수 내부에서 명시적으로 import
        import glob  # 파일 검색을 위한 import
        
        # 보고서 폴더 존재 확인 및 스킵 로직
        docx_folder = "reports_docx.v8"
        md_folder = "generated_reports"
        
        # DOCX 폴더 확인
        docx_files_exist = False
        if os.path.exists(docx_folder):
            docx_files = glob.glob(os.path.join(docx_folder, "*.docx"))
            if docx_files:
                docx_files_exist = True
                print(f"✅ {docx_folder} 폴더에 {len(docx_files)}개의 DOCX 보고서가 이미 존재합니다.")
                for file in docx_files[:3]:  # 최대 3개 파일명만 표시
                    print(f"   - {os.path.basename(file)}")
                if len(docx_files) > 3:
                    print(f"   ... 및 {len(docx_files) - 3}개 더")
        
        # MD 폴더 확인  
        md_files_exist = False
        if os.path.exists(md_folder):
            md_files = glob.glob(os.path.join(md_folder, "*.md"))
            if md_files:
                md_files_exist = True
                print(f"✅ {md_folder} 폴더에 {len(md_files)}개의 MD 보고서가 이미 존재합니다.")
                for file in md_files[:3]:  # 최대 3개 파일명만 표시
                    print(f"   - {os.path.basename(file)}")
                if len(md_files) > 3:
                    print(f"   ... 및 {len(md_files) - 3}개 더")
        
        # 둘 다 존재하면 보고서 생성 과정을 건너뛰기
        if docx_files_exist and md_files_exist:
            print("🚀 기존 보고서가 발견되었습니다. 보고서 생성 과정을 건너뜁니다.")
            return
        
        print("📝 보고서를 새로 생성합니다...")
        load_dotenv()
        # 🌐 LLM 선언
        llm = ChatOpenAI(model="gpt-4o", temperature=0.5)

        # 🏷️ 토픽 제목 생성 프롬프트
        title_prompt_template = ChatPromptTemplate.from_messages([
            ("system",
            "너는 기술 분야 전문가야. 주어진 키워드들을 분석해서 해당 기술 분야를 나타내는 "
            "명확하고 전문적인 제목을 한 문장으로 만들어줘.\n\n"
            "📌 제목 생성 규칙:\n"
            "- 키워드들 간의 기술적 연관성을 파악해서 의미있는 문장으로 구성\n"
            "- 단순 키워드 나열이 아닌, 기술의 목적이나 응용 분야가 드러나는 제목\n"
            "- 10-15자 내외의 간결하면서도 전문적인 표현\n"
            "- '시스템', '기술', '솔루션' 등의 기술 용어 활용 가능\n\n"
            "예시:\n"
            "키워드: 센서, 전력, 온도 → 스마트 센서 기반 전력 효율 제어\n"
            "키워드: 로봇, 주행, 자율 → 자율주행 로봇 네비게이션 시스템\n"
            "키워드: 지능, 인공, 학습 → AI 기반 지능형 학습 플랫폼"
            ),
            ("human",
            "다음 키워드들로 기술 제목을 만들어줘: {keywords}")
        ])

        # 🧾 기술 보고서 작성 프롬프트
        report_prompt_template = ChatPromptTemplate.from_messages([
            ("system",
            "너는 첨단 기술 보고서를 작성하는 전문가이자 기술 전략 컨설턴트야. "
            "다음 주제 제목과 키워드를 기반으로 기업이 실제로 사업화 전략을 세울 수 있도록 기술 보고서를 작성해줘. "
            "보고서 내용에 다시 주제 제목을 명시할 필요는 없어.\n\n"

            "**각 항목은 반드시 '항목명:' 형식으로 시작해.**\n"
            "예: '개요: ...'\n\n"

            "📌 작성 규칙:\n"
            "- 각 항목은 단순 설명이 아니라 기업이 실질적으로 활용 가능한 전략을 포함해야 해.\n"
            "- '기술 구성', '적용 분야', '개발 단계별 목표', '관련 기술 보유 기업 및 제조사 현황' 항목에서는 **각 세부 기술 또는 사례마다 대표 키워드를 번호없이 괄호로 제시한 후 설명하는 방식**으로 작성해줘.\n"
            "  예: (주행 기술) 실내 환경 맵핑과 자율 경로 설정 기술을 통해...\n"
            "- 기술 구성 항목에서는 핵심 기술별 차별화 포인트 및 구현 방안을 제시해.\n"
            "- 적용 분야에서는 수요가 있는 산업 및 시장을 구체적으로 제시하고, 시장 규모나 성장 가능성을 간단히 언급해.\n"
            "- 개발 단계별 목표는 (1차년도), (2차년도), (3차년도) 까지만 제시하고, 연차별 R&D 전략과 실증 또는 제품화 수준을 고려해서 작성해.\n"
            "- 활용 가능성은 기술의 파급력, 확장 가능성, 타 기술과의 융합 가능성을 제안 형태로 서술해.\n"
            "- 관련 기술 보유 기업 및 제조사 현황은 주요 기업별 기술 특징, 차별화 포인트, 벤치마킹 요소를 중심으로 작성해.\n\n"

            "📄 필수 항목:\n"
            "- 개요\n"
            "- 기술 구성\n"
            "- 적용 분야\n"
            "- 개발 단계별 목표\n"
            "- 최종 목표\n"
            "- 활용 가능성\n"
            "- 관련 기술 보유 기업 및 제조사 현황\n\n"

            "보고서는 기업 실무자, 투자자 또는 기술 기획자가 바로 참고할 수 있을 정도로 "
            "**명확하고 전략적이며, 전문적인 문장**으로 구성해줘. "
            "기술 소개가 아니라, 기술 기반 사업 전략 보고서라는 점을 꼭 기억해줘."
            ),

            ("human",
            "다음 정보를 바탕으로 기술 보고서를 작성해줘.\n\n"
            "제목: {title}\n"
            "핵심 키워드: {keywords}")
        ])

        # 📚 섹션 제목 후보 리스트
        SECTION_TITLES = [
            "개요", "기술 구성", "적용 분야",
            "개발 단계별 목표", "최종 목표",
            "활용 가능성", "관련 기술 보유 기업 및 제조사 현황"
        ]

        # 🧹 응답 정리 및 마크다운 Bold 처리
        def clean_and_format_report(report_text):
            paragraphs = []
            current_section = None
            content_accumulator = []

            lines = report_text.strip().split("\n")
            for line in lines:
                line = line.strip()
                if not line:
                    continue

                is_section = False
                for title in SECTION_TITLES:
                    if re.match(fr"^(#+\s*)?\d*[\.\)]?\s*{title}[:：]?$", line, re.IGNORECASE):
                        if current_section:
                            paragraphs.append((current_section, "\n".join(content_accumulator).strip()))
                            content_accumulator = []
                        current_section = title
                        is_section = True
                        break

                if not is_section:
                    content_accumulator.append(line)

            if current_section:
                paragraphs.append((current_section, "\n".join(content_accumulator).strip()))

            return paragraphs

        # 🧾 마크다운 Bold 처리 함수 (텍스트 내 **부분 Bold 처리)
        def add_markdown_paragraph(doc, text):
            p = doc.add_paragraph()

            # 1️⃣ "- 키워드:" 형태를 모두 "(**키워드**)" 로 변경
            text = re.sub(r"^- ([^:]+?)\s*:\s*", r"(\1) ", text)

            # 2️⃣ 마크다운 굵은 텍스트 처리 (**텍스트**)
            pattern = r"(\(\*\*.+?\*\*\))"
            pos = 0
            for match in re.finditer(pattern, text):
                if match.start() > pos:
                    p.add_run(text[pos:match.start()]).font.size = Pt(11)

                bold_text = match.group(2)
                run = p.add_run(bold_text)
                run.bold = True
                run.font.size = Pt(11)
                pos = match.end()

            if pos < len(text):
                p.add_run(text[pos:]).font.size = Pt(11)

            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


        # 📤 Word 저장 함수
        def save_report_as_docx(title, keywords, report_text, filename):
            doc = Document()

            # 제목
            doc.add_heading(title, level=1)

            # 키워드
            p_kw = doc.add_paragraph()
            run_kw = p_kw.add_run(f"핵심 키워드: {keywords}")
            run_kw.bold = True
            run_kw.font.size = Pt(11)

            # 본문 구성
            cleaned_sections = clean_and_format_report(report_text)
            for section_title, content in cleaned_sections:
                # 섹션 제목
                p_title = doc.add_paragraph()
                run_title = p_title.add_run(section_title)
                run_title.bold = True
                run_title.font.size = Pt(13)

                # 섹션 내용
                for paragraph in content.split('\n'):
                    paragraph = paragraph.strip()
                    if paragraph:
                        add_markdown_paragraph(doc, paragraph)

            os.makedirs("reports_docx.v8", exist_ok=True)
            filepath = os.path.join("reports_docx.v8", filename)
            doc.save(filepath)
            print(f"✅ 저장 완료: {filepath}")

        # 🔁 실행 함수 (step4에서 반환되는 딕셔너리 형태 데이터 처리)
        def generate_reports_from_results(topics_dict):
            if not isinstance(topics_dict, dict):
                print("❌ 토픽 데이터가 딕셔너리 형태가 아닙니다.")
                return
                
            for topic_num, keywords in topics_dict.items():
                # 토픽 제목 생성 (LLM이 키워드로 의미있는 문장 생성)
                keywords_for_title = ', '.join(keywords[:5])  # 상위 5개 키워드로 제목 생성
                title_chain = title_prompt_template | llm
                title_response = title_chain.invoke({"keywords": keywords_for_title})
                generated_title = title_response.content.strip() if hasattr(title_response, 'content') else str(title_response).strip()
                
                title = f"Topic {topic_num}: {generated_title}"
                keywords_str = ', '.join(keywords[:10])  # 상위 10개 키워드
                
                print(f"\n🚀 생성 중: {title}")
                report_chain = report_prompt_template | llm
                response = report_chain.invoke({"title": title, "keywords": keywords_str})

                report_text = (
                    response.strip()
                    if isinstance(response, str)
                    else response.content.strip()
                )

                filename = f"{topic_num}_{title[:30].replace(' ', '_').replace(':', '')}.docx"
                save_report_as_docx(title, keywords_str, report_text, filename)

        # ✅ DOCX 보고서 생성 (조건부)
        if not docx_files_exist:
            print("📄 DOCX 보고서를 새로 생성합니다...")
            generate_reports_from_results(x)
        else:
            print("📄 DOCX 보고서가 이미 존재하므로 생성을 건너뜁니다.")
        
        # 표 형태 출력 (딕셔너리를 리스트로 변환)
        table_data = []
        for topic_num, keywords in x.items():
            table_data.append([f"Topic {topic_num}", f"Topic {topic_num}", ', '.join(keywords[:5])])
        
        headers = ["토픽 번호", "토픽", "주요 키워드"]
        print(tabulate(table_data, headers=headers, tablefmt="github"))


        # 보고서 생성 및 저장
        os.makedirs("generated_reports", exist_ok=True)
        
        # MD 보고서가 이미 존재하는지 확인
        if md_files_exist:
            print("📝 MD 보고서가 이미 존재하므로 MD 생성을 건너뜁니다.")
        else:
            print("📝 MD 보고서를 새로 생성합니다...")
            
            # step4에서 반환되는 딕셔너리 형태 데이터 처리
            for topic_id, keywords_list in x.items():
                # 토픽 제목 생성 (LLM이 키워드로 의미있는 문장 생성)
                keywords_for_title = ', '.join(keywords_list[:5])  # 상위 5개 키워드로 제목 생성
                title_chain = title_prompt_template | llm
                title_response = title_chain.invoke({"keywords": keywords_for_title})
                generated_title = title_response.content.strip() if hasattr(title_response, 'content') else str(title_response).strip()
                
                topic_name = f"Topic {topic_id}: {generated_title}"
                keywords_str = ', '.join(keywords_list[:10])  # 상위 10개 키워드를 문자열로 변환
                
                # 1. 보고서 생성
                chain = report_prompt_template | llm
                response = chain.invoke({"title": topic_name, "keywords": keywords_str})
                report_text = response.content.strip()

                # 2. Markdown 파일로 저장
                filename = f"Topic_{topic_id}_{keywords_list[0] if keywords_list else 'empty'}_{keywords_list[1] if len(keywords_list) > 1 else ''}.md"
                filename = filename.replace(' ', '_').replace('/', '_').replace('\\', '_')[:50] + '.md'  # 파일명 정리
                filepath = os.path.join("./generated_reports", filename)

                with open(filepath, "w", encoding="utf-8") as f:
                    f.write(f"# {topic_name}\n")
                    f.write(f"**핵심 키워드:** {keywords_str}\n\n")
                    f.write(report_text)

                print(f"✅ 저장 완료: {filepath}")
